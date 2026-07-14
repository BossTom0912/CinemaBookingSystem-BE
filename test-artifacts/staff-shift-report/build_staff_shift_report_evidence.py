# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
SCREENSHOT = ROOT / "staff_shift_report_screen.png"
DOCX = ROOT / "Staff_Shift_Report_Test_Evidence.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
LIGHT_FILL = "F2F4F7"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf") if bold else Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_card(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, value: str, subtitle: str) -> None:
    draw.rounded_rectangle(box, radius=18, fill="#FFFFFF", outline="#D7DEE8", width=2)
    x1, y1, x2, _ = box
    draw.text((x1 + 26, y1 + 22), title, font=font(24, True), fill="#2A3B4F")
    draw.text((x1 + 26, y1 + 64), value, font=font(46, True), fill="#0B5CAD")
    draw.text((x1 + 26, y1 + 126), subtitle, font=font(20), fill="#5B6775")


def create_report_screen() -> None:
    img = Image.new("RGB", (1600, 1000), "#EEF2F6")
    draw = ImageDraw.Draw(img)

    draw.rectangle((0, 0, 1600, 118), fill="#123A5F")
    draw.text((70, 30), "Staff Shift Report Dashboard", font=font(40, True), fill="#FFFFFF")
    draw.text((70, 78), "GET /api/dashboard/staff/shift-report | Staff A | 2026-07-11 UTC", font=font(22), fill="#DDE9F5")

    draw.rounded_rectangle((70, 155, 1530, 245), radius=18, fill="#FFFFFF", outline="#D7DEE8", width=2)
    draw.text((100, 178), "Scope", font=font(23, True), fill="#2A3B4F")
    draw.text((210, 178), "Staff self report only", font=font(23), fill="#2A3B4F")
    draw.text((100, 212), "Cinema", font=font(23, True), fill="#2A3B4F")
    draw.text((210, 212), "CIN_SHIFT_A - Shift Cinema A", font=font(23), fill="#2A3B4F")
    draw.text((820, 178), "Auth Policy", font=font(23, True), fill="#2A3B4F")
    draw.text((985, 178), "CanViewStaffShiftReport", font=font(23), fill="#2A3B4F")
    draw.text((820, 212), "Source", font=font(23, True), fill="#2A3B4F")
    draw.text((985, 212), "Integration test fixture data", font=font(23), fill="#2A3B4F")

    draw_card(draw, (70, 285, 430, 470), "Tickets checked in", "2", "CHECKIN_LOG result SUCCESS")
    draw_card(draw, (460, 285, 820, 470), "F&B fulfilled", "3", "2 counter + 1 online")
    draw_card(draw, (850, 285, 1210, 470), "Counter revenue", "200", "120 cash + 80 transfer")
    draw_card(draw, (1240, 285, 1530, 470), "Transactions", "5", "Rows returned in detail list")

    headers = ["Time", "Type", "Reference", "Amount", "Payment"]
    rows = [
        ["14:00", "ONLINE_FB_FULFILLMENT", "BKG_SHIFT_A_ONLINE_FB", "80", "-"],
        ["10:30", "COUNTER_FB_ORDER", "BKG_SHIFT_A_TRANSFER", "80", "BANK_TRANSFER"],
        ["10:00", "TICKET_CHECK_IN", "CIL_SHIFT_A_2", "0", "-"],
        ["09:30", "COUNTER_FB_ORDER", "BKG_SHIFT_A_CASH", "120", "CASH"],
        ["09:00", "TICKET_CHECK_IN", "CIL_SHIFT_A_1", "0", "-"],
    ]
    table_x, table_y = 70, 530
    col_w = [180, 390, 420, 150, 250]
    row_h = 62
    draw.rounded_rectangle((table_x, table_y, 1530, table_y + row_h * (len(rows) + 1)), radius=14, fill="#FFFFFF", outline="#D7DEE8", width=2)
    x = table_x
    for idx, header in enumerate(headers):
        draw.rectangle((x, table_y, x + col_w[idx], table_y + row_h), fill="#E8EEF5")
        draw.text((x + 18, table_y + 18), header, font=font(21, True), fill="#1F3A5F")
        x += col_w[idx]
    for r_idx, row in enumerate(rows):
        y = table_y + row_h * (r_idx + 1)
        x = table_x
        for c_idx, value in enumerate(row):
            draw.line((x, y, x + col_w[c_idx], y), fill="#D7DEE8", width=1)
            draw.text((x + 18, y + 18), value, font=font(20), fill="#263545")
            x += col_w[c_idx]

    draw.text((70, 925), "Note: Screenshot generated from the passing API integration-test dataset; no frontend screen exists yet.", font=font(22), fill="#5B6775")
    img.save(SCREENSHOT)


def set_font(run, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_col_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths: list[float], header: bool = True) -> None:
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_col_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_font(run, 9.5, RGBColor(0, 0, 0), row_idx == 0 and header)
            if row_idx == 0 and header:
                shade_cell(cell, LIGHT_FILL)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 13, BLUE if level == 1 else DARK_BLUE, True)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, 11, RGBColor(0, 0, 0), True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, 11, RGBColor(0, 0, 0))
    else:
        run = p.add_run(text)
        set_font(run, 11, RGBColor(0, 0, 0))


def create_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.text = "CinemaSystem - Staff Shift Report Evidence"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = GRAY

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Staff Shift Report Dashboard")
    set_font(r, 24, RGBColor(0, 0, 0), True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("Test evidence and report screenshot for GET /api/dashboard/staff/shift-report")
    set_font(r, 12, GRAY)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    values = [
        ("Branch", "Tom/staff-shift-report-dashboard"),
        ("Test command", "dotnet test CinemaSystem.Tests\\CinemaSystem.Tests.csproj --no-build --filter StaffShiftReportApiIntegrationTests"),
        ("Result", "Passed: 8, Failed: 0, Skipped: 0"),
        ("Generated", "2026-07-11"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
    style_table(meta, [1.55, 4.95], header=False)

    add_heading(doc, "Evidence Summary", 1)
    add_body(doc, "Outcome: targeted Staff Shift Report integration tests passed 8/8.", "Outcome:")
    add_body(doc, "Scope rule: Staff can see self only; Manager is limited to assigned cinema; Admin can bypass cinema scope; Customer is forbidden.", "Scope rule:")
    add_body(doc, "Data rule: report uses CHECKIN_LOG success rows, BOOKING counter orders, and BOOKING online F&B fulfillment rows tracked by fbFulfilledByStaffProfileId.", "Data rule:")

    add_heading(doc, "Screenshot", 1)
    add_body(doc, "The following screen is generated from the passing integration-test fixture data because this task exposes a backend API and no frontend page exists yet.")
    doc.add_picture(str(SCREENSHOT), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run("Figure 1. Staff Shift Report dashboard evidence screen")
    set_font(r, 9.5, GRAY)

    add_heading(doc, "Functional Test Matrix", 1)
    matrix = doc.add_table(rows=1, cols=4)
    matrix.style = "Table Grid"
    headers = ["Case", "Actor", "Expected", "Status"]
    for idx, h in enumerate(headers):
        matrix.rows[0].cells[idx].text = h
    rows = [
        ("Staff xem báo cáo của chính mình", "Staff", "200 OK", "Passed"),
        ("Staff xem báo cáo Staff khác", "Staff", "403 Forbidden", "Passed"),
        ("Manager xem Staff cùng cinema", "Manager", "200 OK", "Passed"),
        ("Manager xem Staff khác cinema", "Manager", "403 Forbidden", "Passed"),
        ("Customer gọi endpoint", "Customer", "403 Forbidden", "Passed"),
        ("Admin gọi không filter", "Admin", "200 OK", "Passed"),
        ("Date range sai hoặc thiếu tham số", "Admin", "400 Bad Request", "Passed"),
        ("Calculation validation", "Staff/Admin", "Counts and revenue match fixture", "Passed"),
    ]
    for item in rows:
        cells = matrix.add_row().cells
        for idx, value in enumerate(item):
            cells[idx].text = value
    style_table(matrix, [2.35, 1.0, 1.85, 1.3])

    add_heading(doc, "Sample Report Metrics", 1)
    metrics = doc.add_table(rows=1, cols=3)
    metrics.style = "Table Grid"
    for idx, h in enumerate(["Metric", "Value", "Source"]):
        metrics.rows[0].cells[idx].text = h
    metric_rows = [
        ("Checked-in tickets", "2", "CHECKIN_LOG where Result = SUCCESS"),
        ("Fulfilled F&B orders", "3", "2 counter + 1 online fulfillment"),
        ("Counter revenue", "200", "BOOKING totalAmount for COUNTER orders"),
        ("Cash revenue", "120", "PaymentMethod = CASH"),
        ("Transfer revenue", "80", "PaymentMethod = BANK_TRANSFER"),
        ("Transaction detail rows", "5", "Combined report transactions"),
    ]
    for item in metric_rows:
        cells = metrics.add_row().cells
        for idx, value in enumerate(item):
            cells[idx].text = value
    style_table(metrics, [2.0, 1.0, 3.5])

    doc.save(DOCX)


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    create_report_screen()
    create_docx()
    print(DOCX)
    print(SCREENSHOT)


if __name__ == "__main__":
    main()
